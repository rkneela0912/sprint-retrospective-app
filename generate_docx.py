"""Generate blog-article.docx using python-docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import copy

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

# ── Helper: set paragraph border (bottom) ─────────────────────────────────────
def add_bottom_border(paragraph, color="E8ECF2"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Helper: set paragraph background shading ──────────────────────────────────
def shade_paragraph(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page margins (2 cm all round)
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Base styles ───────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Segoe UI"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Heading 1
h1 = doc.styles["Heading 1"]
h1.font.name = "Segoe UI"
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

# Heading 2
h2 = doc.styles["Heading 2"]
h2.font.name = "Segoe UI"
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(4)

# Heading 3
h3 = doc.styles["Heading 3"]
h3.font.name = "Segoe UI"
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(2)

# ── TITLE / HERO ──────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
shade_paragraph(title_p, "F0F7FF")
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(0)
title_p.paragraph_format.left_indent = Inches(0.3)
run = title_p.add_run("POWER PLATFORM · REACT · SHAREPOINT")
run.font.name = "Segoe UI"
run.font.size = Pt(8)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)

title_h = doc.add_heading(
    "Building a Sprint Retrospective App on Power Apps Code Apps", level=1
)
title_h.paragraph_format.space_before = Pt(6)
title_h.paragraph_format.space_after = Pt(4)
shade_paragraph(title_h, "F0F7FF")
title_h.paragraph_format.left_indent = Inches(0.3)

meta_p = doc.add_paragraph()
shade_paragraph(meta_p, "F0F7FF")
meta_p.paragraph_format.left_indent = Inches(0.3)
meta_p.paragraph_format.space_after = Pt(14)
run = meta_p.add_run("March 2026  ·  ~12 min read  ·  Architecture · SharePoint · Lessons Learned")
run.font.name = "Segoe UI"
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x6B, 0x7C, 0x93)


def h2(text):
    p = doc.add_heading(text, level=2)
    add_bottom_border(p, "CBD5E1")
    return p


def h3(text):
    return doc.add_heading(text, level=3)


def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)
        p.add_run(text).font.name = "Segoe UI"
    else:
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(11)
    return p


def callout(text):
    p = doc.add_paragraph()
    shade_paragraph(p, "F0F7FF")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x2C, 0x40, 0x60)
    return p


def code_block(text):
    p = doc.add_paragraph()
    shade_paragraph(p, "0D1117")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Cascadia Code"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xE6, 0xED, 0xF3)
    return p


def insert_screenshot(img_path, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption_text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x6B, 0x7C, 0x93)


def simple_table(headers, rows, header_color="0078D4"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, header_color)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row in enumerate(rows):
        trow = t.rows[ri + 1]
        bg = "F8FBFF" if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            shade_cell(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
    doc.add_paragraph()  # spacer


# ── SECTION 1: THE PROBLEM ────────────────────────────────────────────────────
h2("The Problem: Retrospectives Are Broken")
body(
    "Most agile teams know the value of a good sprint retrospective. In theory it is a structured conversation "
    "about what went well, what did not, and what needs to change. In practice it often looks like this: someone "
    "shares a Miro board or a sticky-note Jam, the team fills it in, someone reads the notes aloud, and action "
    "items get written into a chat message that nobody looks at again."
)
body("Three things consistently go wrong:")
numbered(
    "Action items fall through the cracks. There is no formal ownership, no status tracking, and no connection "
    "between an action and the feedback that generated it."
)
numbered(
    "Psychological safety is fragile. When feedback is always attributable, team members self-censor. "
    "The most important things often go unsaid."
)
numbered(
    "There is no continuity between sprints. Incomplete actions from Sprint 5 quietly disappear when Sprint 6 starts."
)
body(
    "The goal of this project was to solve all three inside the Microsoft 365 ecosystem — no new SaaS "
    "subscriptions, no external databases, no change-management headaches."
)

# ── SECTION 2: WHY POWER APPS CODE APPS ──────────────────────────────────────
h2("Why Power Apps Code Apps?")
body(
    "Power Apps Code Apps is a relatively new capability in the Power Platform that lets you build a Power App "
    "using a standard web front-end stack — React, TypeScript, Vite — and deploy it as a first-class Power App. "
    "You get:"
)
bullet("Microsoft Entra (Azure AD) authentication with zero configuration")
bullet("SharePoint Online connectors with auto-generated TypeScript service clients")
bullet("Deployment via a single CLI command (pac code push)")
bullet("The app lives in your Power Apps environment alongside canvas and model-driven apps")
body(
    "The alternative was a standalone Azure Static Web App with its own authentication and SharePoint Graph API "
    "integration. That is valid, but it adds infrastructure, a separate identity configuration, and a longer path "
    "from 'developer laptop' to 'available to the team in Teams.'"
)
body(
    "For an internal productivity tool used by a single organisation, Power Apps Code Apps is the right trade-off. "
    "You get modern web development tooling while staying firmly inside the Microsoft 365 governance boundary."
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\Power Apps maker portal showing the deployed Sprint Retrospective app alongside canvas apps in the same environment..png",
    "Power Apps maker portal showing the deployed Sprint Retrospective app alongside canvas apps in the same environment.",
)

# ── SECTION 3: WHAT THE APP DOES ─────────────────────────────────────────────
h2("What the App Does")
body("The app is a four-column retrospective board with a structured session lifecycle.")

h3("The Four Columns")
simple_table(
    ["Column", "Purpose"],
    [
        ["What We Liked", "Celebrate what worked well this sprint"],
        ["What We Missed", "Honest reflection on gaps and problems"],
        ["What We Learned", "Knowledge gained — technical, process, team"],
        ["Appreciations", "Shout-outs and recognition between team members"],
    ],
)

h3("Session Lifecycle")
body("Every retrospective is a session that moves through three states:")
lifecycle = doc.add_paragraph()
lifecycle.paragraph_format.space_after = Pt(8)
for label, sep in [("Draft", " → "), ("Open", " → "), ("Closed", "")]:
    r = lifecycle.add_run(label)
    r.font.name = "Segoe UI"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    if sep:
        s = lifecycle.add_run(sep)
        s.font.name = "Segoe UI"
        s.font.size = Pt(11)
        s.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)

bullet("Draft — Session created and configured by the Scrum Master. No submissions accepted yet.")
bullet("Open — The board is live. Team members can add items, vote, and attach action items.")
bullet("Closed — Board is read-only. Scrum Master can carry incomplete actions over to the next sprint.")

h3("Role-Based Access")
simple_table(
    ["Action", "Scrum Master", "Team Member", "Viewer"],
    [
        ["View retro board", "✅", "✅", "✅"],
        ["Create sessions", "✅", "❌", "❌"],
        ["Change session status", "✅", "❌", "❌"],
        ["Add retro items", "✅", "✅", "❌"],
        ["Vote on items", "✅", "✅", "❌"],
        ["Add action items", "✅", "✅", "❌"],
        ["Edit action item status", "✅", "Owner only", "❌"],
        ["Carry over actions", "✅", "❌", "❌"],
        ["See anonymous authors", "✅", "❌", "❌"],
    ],
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\The Sprint Selector view showing the list of retrospective sessions with their status badges (Draft  Open  Closed)..png",
    "The Sprint Selector view showing the list of retrospective sessions with their status badges (Draft / Open / Closed).",
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\The main retrospective board with all four columns, feedback cards, and vote counts visible. Session status shown in the header..png",
    "The main retrospective board with all four columns, feedback cards, and vote counts visible. Session status shown in the header.",
)

# ── SECTION 4: ARCHITECTURE ───────────────────────────────────────────────────
h2("Architecture")

h3("Tech Stack")
simple_table(
    ["Layer", "Choice", "Rationale"],
    [
        ["UI Framework", "React 19 + TypeScript", "Standard modern web stack; strong ecosystem"],
        ["Build Tool", "Vite 7", "Fast HMR in dev, optimised production bundle"],
        ["Platform SDK", "@microsoft/power-apps", "Code Apps runtime context (user identity, connectors)"],
        ["Data Backend", "SharePoint Online", "Already available in every Microsoft 365 tenant"],
        ["Auth", "Power Apps getContext()", "Zero-config Entra ID auth"],
        ["Deployment", "pac code push", "Single command deploy to Power Apps environment"],
    ],
)

h3("The Data Model — Four SharePoint Lists")
body(
    "All data lives in SharePoint Online. No Azure SQL, no Dataverse, no Cosmos DB. This is a deliberate choice "
    "for organisations that want simplicity and zero extra licensing cost."
)
body("RetroSessions — one record per retrospective session")
code_block(
    "Title        | Single line of text  | Session name\n"
    "Sprint       | Choice               | Sprint 81, Sprint 82, …\n"
    "Status       | Choice               | Draft | Open | Closed\n"
    "SessionDate  | Date and Time        |"
)
body("RetroItems — one record per piece of feedback")
code_block(
    "Title        | Single line of text  | The feedback text\n"
    "Category     | Choice               | What We Liked | What We Missed | What We Learned | Appreciations\n"
    "SessionRef   | Number               | FK → RetroSessions.ID\n"
    "SubmittedBy  | Person or Group      |\n"
    "Votes        | Number               | Default: 0\n"
    "IsAnonymous  | Yes/No               | Default: No"
)
body("ActionItems — one record per action item")
code_block(
    "Title           | Single line of text  | Action description\n"
    "Owner           | Person or Group      |\n"
    "Status          | Choice               | Open | In Progress | Completed | Carried Over\n"
    "TargetSprint    | Choice               | Sprint 81, Sprint 82, …\n"
    "RetroItemRef    | Number               | FK → RetroItems.ID\n"
    "CarriedOverFrom | Number               | FK → ActionItems.ID (null if original)"
)
body("UserRoles — one record per user")
code_block(
    "Title  | Single line of text  | User's email / UPN\n"
    "Role   | Choice               | Scrum Master | Team Member | Viewer"
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\The four SharePoint lists in the SharePoint site's list view..png",
    "The four SharePoint lists in the SharePoint site's list view.",
)

# ── SECTION 5: KEY TECHNICAL DECISIONS ───────────────────────────────────────
h2("Key Technical Decisions")

h3("1. Identity and Role Resolution")
body(
    "Power Apps Code Apps provide a getContext() function from the @microsoft/power-apps SDK. "
    "This returns the signed-in user's userPrincipalName — used as the key into the UserRoles SharePoint list."
)
code_block(
    "// src/components/AuthProvider.tsx\n"
    "import { getContext } from '@microsoft/power-apps/lib/app';\n\n"
    "const ctx = await getContext();\n"
    "const email = ctx?.user?.userPrincipalName?.toLowerCase() ?? '';\n\n"
    "const result = await UserRolesService.getAll({\n"
    "  filter: `Email eq '${email}' and IsActive eq 1`,\n"
    "  top: 1,\n"
    "});\n\n"
    "const role = result.data[0]?.Role?.Value ?? 'Team Member';\n"
    "setUserRole(role as UserRole);"
)
callout(
    "Design decision: Users not in the UserRoles list default to 'Team Member' rather than 'Viewer' — "
    "a permissive default that reduces onboarding friction. Change this fallback in AuthProvider.tsx "
    "if you need a restrictive default."
)

h3("2. Anonymous Submissions Without Anonymous Storage")
body(
    "The IsAnonymous flag is stored on the RetroItem record alongside the actual SubmittedBy person field. "
    "Anonymity is enforced by the application layer, not the data layer."
)
body(
    "The practical consequence: the Scrum Master can always see who submitted anonymously, displayed as "
    "'🔒 Anonymous (user@domain.com)' — visible only to them. This provides a safety net against abuse "
    "while preserving psychological safety for team members."
)
code_block(
    "// src/components/RetroCard.tsx\n"
    "const displayName =\n"
    "  item.isAnonymous && userRole !== 'Scrum Master'\n"
    "    ? 'Anonymous'\n"
    "    : item.submittedBy;"
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\A feedback card showing the anonymous badge as seen by a Team Member (left) vs. the Scrum Master view showing the author in the badge (right)..png",
    "A feedback card showing the anonymous badge as seen by a Team Member (left) vs. the Scrum Master view showing the author in the badge (right).",
)

h3("3. Optimistic Voting")
body(
    "The vote button uses optimistic UI: the local vote count increments immediately before the SharePoint "
    "write confirms. If the write fails, the count rolls back. The service call passes the explicit nextVotes "
    "value rather than incrementing server-side — avoiding the race condition where two concurrent votes "
    "silently overwrite each other."
)
code_block(
    "// src/components/RetroCard.tsx\n"
    "const handleVote = async () => {\n"
    "  if (isVoting || sessionStatus === 'Closed') return;\n"
    "  setIsVoting(true);\n"
    "  const nextVotes = localVotes + 1;\n"
    "  setLocalVotes(nextVotes); // immediate optimistic update\n"
    "  try {\n"
    "    await voteOnItem(item.id, nextVotes);\n"
    "  } catch (err) {\n"
    "    setLocalVotes((prev) => Math.max(prev - 1, 0)); // rollback\n"
    "  } finally {\n"
    "    setIsVoting(false);\n"
    "  }\n"
    "};"
)

h3("4. Carry-Over Logic")
body(
    "This is the feature that solves 'action items fall through the cracks.' When the Scrum Master closes a "
    "sprint and opens the next one, they can run the carry-over from the Admin Panel. The logic queries all "
    "Open or In Progress action items from the source session and creates new copies in the target session, "
    "with a CarriedOverFrom foreign key pointing back to the original."
)
code_block(
    "// src/services/SharePointService.ts\n"
    "export const carryOverActions = async (\n"
    "  fromSessionId: number,\n"
    "  toSessionId: number,\n"
    "  toSprint: string\n"
    "): Promise<void> => {\n"
    "  const result = await ActionItemsService.getAll({\n"
    "    filter: `SessionRefID eq ${fromSessionId} and\n"
    "             (Status eq 'Open' or Status eq 'In Progress')`,\n"
    "  });\n"
    "  const carryOverPromises = (result.data as any[]).map((action) =>\n"
    "    ActionItemsService.create({\n"
    "      Title: action.Title,\n"
    "      Owner: action.Owner,\n"
    "      SessionRefID: toSessionId,\n"
    "      Status: { Value: 'Carried Over' },\n"
    "      TargetSprint: { Value: toSprint },\n"
    "      CarriedOverFrom: action.ID,\n"
    "    })\n"
    "  );\n"
    "  await Promise.all(carryOverPromises);\n"
    "};"
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\The Admin Panel slide-out (Scrum Master only) showing the session status controls and the carry-over sprint selector..png",
    "The Admin Panel slide-out (Scrum Master only) showing the session status controls and the carry-over sprint selector.",
)

h3("5. SharePoint People Field Resolution")
body(
    "Writing to a SharePoint Person or Group column via the Power Apps connector requires a properly shaped "
    "object, not just an email string. The resolvePersonValue helper calls the connector's "
    "getReferencedEntity operation to look up the user and returns either an exact UPN match or a "
    "Claims-based fallback object."
)
code_block(
    "// src/services/SharePointService.ts\n"
    "async function resolvePersonValue(search, referencedKey, service) {\n"
    "  const res = await service.getReferencedEntity(search, referencedKey);\n"
    "  const candidates = extractCandidates(res);\n"
    "  if (candidates.length > 0) {\n"
    "    const exact =\n"
    "      candidates.find((c) => toLower(c?.Email) === needle) ||\n"
    "      candidates.find((c) => toLower(c?.UserPrincipalName) === needle);\n"
    "    return exact ?? candidates[0];\n"
    "  }\n"
    "  // Claims fallback\n"
    "  return {\n"
    "    Claims: `i:0#.f|membership|${needle}`,\n"
    "    Email: needle,\n"
    "    DisplayName: needle,\n"
    "  };\n"
    "}"
)
body(
    "This is one of the trickiest parts of the Power Apps Code Apps connector model. The auto-generated "
    "clients handle the HTTP plumbing, but writing Person and Choice fields requires knowing the exact "
    "shape the connector expects. The claims fallback covers the majority of tenant configurations."
)

# ── SECTION 6: PROJECT STRUCTURE ─────────────────────────────────────────────
h2("Project Structure")
code_block(
    "sprint-retro-app/\n"
    "├── src/\n"
    "│   ├── components/\n"
    "│   │   ├── AuthProvider.tsx        # User identity & role resolution\n"
    "│   │   ├── SprintSelectorView.tsx  # Session list / create session\n"
    "│   │   ├── RetrospectiveView.tsx   # Main board page\n"
    "│   │   ├── RetroBoard.tsx          # Four-column board layout\n"
    "│   │   ├── RetroCard.tsx           # Feedback card with votes & actions\n"
    "│   │   ├── AdminPanel.tsx          # Scrum Master admin slide-out\n"
    "│   │   └── …\n"
    "│   ├── services/                   # Business logic wrapping generated clients\n"
    "│   │   └── SharePointService.ts    # Core data operations\n"
    "│   └── generated/                  # Auto-generated by pac CLI — do not edit\n"
    "├── power.config.json               # Power Apps deployment configuration\n"
    "└── vite.config.ts"
)
body(
    "The generated/ folder contains service clients produced by the Power Platform CLI when a SharePoint data "
    "source is added. The services/ layer wraps them with typed business operations and normalisation helpers, "
    "keeping the UI components clean and testable."
)

# ── SECTION 7: DEPLOYMENT ─────────────────────────────────────────────────────
h2("Deployment")
body(
    "Deploying to Power Apps is a three-command workflow once the SharePoint lists exist and the "
    "Power Platform CLI is authenticated:"
)
code_block(
    "# Authenticate against your Power Apps environment\n"
    "pac auth create --environment <your-environment-id>\n\n"
    "# Build the React app\n"
    "npm run build\n\n"
    "# Push to Power Apps\n"
    "pac code push"
)
body(
    "The power.config.json file maps SharePoint connector connection IDs and list URLs to the generated "
    "service names. First-time setup also requires running pac code add-data-source once per SharePoint "
    "list to generate the service clients and register the connector."
)
insert_screenshot(
    r"C:\Apps\sprint-retro-app\Screenshots\pac-code-push-terminal.png",
    "Terminal output of 'pac code push' showing the successful deployment and the app URL in the Power Apps environment.",
)

# ── SECTION 8: WHAT WE'D DO DIFFERENTLY ──────────────────────────────────────
h2("What We Would Do Differently")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("1. Dataverse instead of SharePoint for a production deployment at scale.  ")
r.bold = True
r.font.name = "Segoe UI"
r.font.size = Pt(11)
p.add_run(
    "SharePoint lists work well and have zero extra cost, but Dataverse gives you proper relational integrity, "
    "server-side business rules, and better Power Automate integration. For a team of 10–15 people running "
    "weekly retros, SharePoint is fine. For a multi-team deployment, Dataverse is the right foundation."
).font.name = "Segoe UI"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("2. Power Automate for notifications.  ")
r.bold = True
r.font.name = "Segoe UI"
r.font.size = Pt(11)
p.add_run(
    "The current app has no notifications. A simple Power Automate flow triggered on an ActionItem status "
    "change — 'Your action item X was updated to In Progress' — would close the feedback loop without any "
    "changes to the React codebase."
).font.name = "Segoe UI"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("3. Sprint choices driven by data, not hardcoded constants.  ")
r.bold = True
r.font.name = "Segoe UI"
r.font.size = Pt(11)
p.add_run(
    "Sprint labels are currently hardcoded in RetroCard.tsx. They should come from the RetroSessions list "
    "or a configuration list so the app does not need a code change when the sprint numbering scheme changes."
).font.name = "Segoe UI"

# ── SECTION 9: CLOSING THOUGHTS ───────────────────────────────────────────────
h2("Closing Thoughts")
body(
    "Power Apps Code Apps is still a relatively young capability in the Power Platform, but it already closes "
    "a significant gap: you can now bring real software engineering practices — React, TypeScript, component "
    "architecture, CI/CD — to Power Apps without sacrificing the zero-friction authentication and connector "
    "ecosystem that makes the platform valuable to Microsoft 365 organisations."
)
body(
    "For this sprint retrospective use case, the combination of Code Apps, SharePoint Online, and a handful "
    "of well-designed React components produced a tool that is genuinely useful — anonymous submissions, "
    "action tracking with carry-over, role-based permissions — in a way that fits naturally inside the tools "
    "a team is already using."
)
body(
    "The architecture is intentionally simple. It could be made more robust with Dataverse, more observable "
    "with Application Insights, and more integrated with Power Automate notifications. But as a starting "
    "point for any team frustrated by ad-hoc sticky-note retros, it works — and it can be stood up in an afternoon."
)

# ── CTA ────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
cta = doc.add_paragraph()
shade_paragraph(cta, "0078D4")
cta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
cta.paragraph_format.space_before = Pt(12)
cta.paragraph_format.space_after = Pt(0)
r = cta.add_run("Want to Build Something Similar?")
r.font.name = "Segoe UI"
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

cta2 = doc.add_paragraph()
shade_paragraph(cta2, "0078D4")
cta2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
cta2.paragraph_format.space_before = Pt(4)
cta2.paragraph_format.space_after = Pt(0)
r = cta2.add_run(
    "If you are exploring Power Apps Code Apps for your own internal tools, or looking for a development "
    "partner to build productivity applications on the Microsoft 365 platform, I would love to have a conversation."
)
r.font.name = "Segoe UI"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0xE0, 0xF0, 0xFF)

cta3 = doc.add_paragraph()
shade_paragraph(cta3, "0078D4")
cta3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
cta3.paragraph_format.space_before = Pt(8)
cta3.paragraph_format.space_after = Pt(16)
r = cta3.add_run("Ranjith Neela\n")
r.font.name = "Segoe UI"
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
r2 = cta3.add_run("iamranjithneela@gmail.com  ·  www.ranjithneela.com")
r2.font.name = "Segoe UI"
r2.font.size = Pt(10)
r2.font.bold = False
r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ── FOOTER ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run(
    "Built with React 19 · TypeScript · Vite 7 · Power Apps Code Apps SDK · SharePoint Online\n"
    "Deployed via Power Platform CLI (pac code push)"
)
r.font.name = "Segoe UI"
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save("blog-article.docx")
print("blog-article.docx created successfully.")
